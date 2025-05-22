
# main.py
import os
import base64
import io
import math
from flask import Flask, render_template, Response, redirect, request, session, abort, url_for
from flask_mail import Mail, Message
from flask import send_file
import mysql.connector
import hashlib
import datetime
from datetime import datetime
from datetime import date
import calendar
import random
from random import randint
from urllib.request import urlopen
import webbrowser
####
import docx
from docx import Document
from pdf2docx import parse
from typing import Tuple
from PIL import Image, ImageDraw, ImageFont
import textwrap

from spire.doc import *
from spire.doc.common import *
###

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

from werkzeug.utils import secure_filename

import urllib.request
import urllib.parse
import socket    
import csv

import matplotlib as mpl
import seaborn as sns
from matplotlib import pyplot as plt
from collections import OrderedDict

import re    # for regular expressions 
#import nltk  # for text manipulation 
import string # for text manipulation 
import warnings
#from nltk.stem.porter import *
#from nltk.corpus import stopwords
#nltk.download()
#from nltk.stem import PorterStemmer
#from nltk.tokenize import sent_tokenize, word_tokenize
#from nltk.corpus import stopwords
###

import gensim
from gensim.parsing.preprocessing import remove_stopwords, STOPWORDS
from gensim.parsing.porter import PorterStemmer

#from spacy.lang.es import Spanish
#nlp = Spanish()

import email.policy
from bs4 import BeautifulSoup
#import tensorflow as tf
#from tensorflow.keras.preprocessing.text import Tokenizer
#from tensorflow.keras.preprocessing.sequence import pad_sequences
from sklearn.metrics import confusion_matrix
from sklearn.model_selection import train_test_split
import sklearn
###

#from imblearn.over_sampling import SMOTEN
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.metrics import classification_report, plot_confusion_matrix
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import MultinomialNB
from sklearn.pipeline import make_pipeline
from wordcloud import WordCloud
from collections import Counter

plt.rc("axes.spines", right=False, top=False)
plt.rc("font", family="serif")
###
'''from tqdm import tqdm
from keras.preprocessing.text import Tokenizer
tqdm.pandas(desc="progress-bar")
#from gensim.models import Doc2Vec
from sklearn import utils
from sklearn.model_selection import train_test_split
from keras.preprocessing.sequence import pad_sequences

from sklearn.linear_model import LogisticRegression
#from gensim.models.doc2vec import TaggedDocument
import re'''
###
###########
stemmer = PorterStemmer()
    
from wordcloud import STOPWORDS
STOPWORDS.update(['rt', 'mkr', 'didn', 'bc', 'n', 'm', 
                  'im', 'll', 'y', 've', 'u', 'ur', 'don', 
                  'p', 't', 's', 'aren', 'kp', 'o', 'kat', 
                  'de', 're', 'amp', 'will'])

###########

mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  password="",
  charset="utf8",
  database="malicious_email"

)
app = Flask(__name__)
##session key
app.secret_key = 'abcdef'
#######
UPLOAD_FOLDER = 'static/upload'
ALLOWED_EXTENSIONS = { 'csv'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
#####
##email
mail_settings = {
    "MAIL_SERVER": 'smtp.gmail.com',
    "MAIL_PORT": 465,
    "MAIL_USE_TLS": False,
    "MAIL_USE_SSL": True,
    "MAIL_USERNAME": "rndittrichy@gmail.com",
    "MAIL_PASSWORD": "lyylfimewwddjwnk"
}

app.config.update(mail_settings)
mail = Mail(app)
#######

def sendmail(usermail,mess1,fid,n1):

    subj1="Mail-Alert"
    with app.app_context():
        msg = Message(subject=subj1, sender=app.config.get("MAIL_USERNAME"),recipients=[usermail], body=mess1)
        if n1>0:
            r=1
                                                        
            while r<=n1:
                sm="m"+str(fid)+"_"+str(r)+".png"
                with app.open_resource("static/attachments/"+sm) as fp2:
                    msg.attach("static/attachments/"+sm, "images/png", fp2.read())
                r+=1
        mail.send(msg)

@app.route('/', methods=['GET', 'POST'])
def index():

    '''txt="hai hello how are you in the program"
    t=txt.split(" ")
    #print("aaa")
    #print(t)
    stop_words = ['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've", "you'll", "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', "wouldn't",',','.','I','\'','-','/']
    ##Tokenize
    data1=[]
    i=0
    for ds in t:
        dt=[]
        if i<5:
            
            dt.append(ds[1])
            text=ds[1]

            doc = nlp(text)
            text_tokens = [token.text for token in doc]
            #text_tokens=tokenize_by_word(text)
            #text_tokens =word_tokenize(text)
            tokens_without_sw = [word for word in text_tokens if not word in stop_words]
            dt.append(tokens_without_sw)

            
            data1.append(dt)
        i+=1
    print(data1)'''

    return render_template('index.html')

    
@app.route('/login_user', methods=['GET', 'POST'])
def login_user():
    msg=""

    act=request.args.get("act")
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM register WHERE uname = %s AND pass = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
            ff=open("user.txt","w")
            ff.write(uname)
            ff.close()
            return redirect(url_for('userhome'))
        else:
            msg = 'Incorrect username/password!'

    
    return render_template('login_user.html',msg=msg,act=act)


@app.route('/login', methods=['GET', 'POST'])
def login():
    msg=""
    act=request.args.get("act")
    #usermail=""
    #mess1="mytest"
    #sendmail(usermail,mess1)
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM admin WHERE username = %s AND password = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
            return redirect(url_for('admin'))
        else:
            msg = 'Incorrect username/password!'
    return render_template('login.html',msg=msg,act=act)

@app.route('/register', methods=['GET', 'POST'])
def register():
   
    msg=""
    act=request.args.get("act")
    if request.method=='POST':
        name=request.form['name']
        mobile=request.form['mobile']
        #email=request.form['email']
        uname=request.form['uname']
        pass1=request.form['pass']
        #password=request.form['password']
       
        mycursor = mydb.cursor()
        mycursor.execute("SELECT count(*) FROM register where uname=%s",(uname,))
        cnt = mycursor.fetchone()[0]

        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM register")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
                    
            sql = "INSERT INTO register(id,name,mobile,uname,pass) VALUES (%s, %s, %s, %s, %s)"
            val = (maxid,name,mobile,uname,pass1)
            mycursor.execute(sql, val)
            mydb.commit()            
            #print(mycursor.rowcount, "Registered Success")
            msg="success"
            #if mycursor.rowcount==1:
            return redirect(url_for('register',act='1'))
        else:
            msg='fail'
    return render_template('register.html',msg=msg,act=act)

def convert_pdf2docx(input_file: str, output_file: str, pages: Tuple = None):
    """Converts pdf to docx"""
    if pages:
        pages = [int(i) for i in list(pages) if i.isnumeric()]
    result = parse(pdf_file=input_file,
                   docx_with_path=output_file, pages=pages)
    summary = {
        "File": input_file, "Pages": str(pages), "Output File": output_file
    }
    # Printing Summary
    print("## Summary ########################################################")
    print("\n".join("{}:{}".format(i, j) for i, j in summary.items()))
    print("###################################################################")
    return result

def text_to_image(text_file, fid):

    image_file="m"+str(fid)+"_1.png"
    font_path="static/arial.ttf"
    font_size=20
    wrap_width=80
    with open("static/attachments/"+text_file, 'r') as f:
        text = f.read()

    # Wrap the text
    wrapped_text = textwrap.wrap(text, width=wrap_width)

    # Create a new image with appropriate dimensions
    line_height = font_size + 5  # Add some spacing between lines
    img_width = max(len(line) * font_size for line in wrapped_text) * 2
    img_height = len(wrapped_text) * line_height
    img = Image.new("RGB", (img_width, img_height), "white")
    draw = ImageDraw.Draw(img)

    # Load the font
    font = ImageFont.truetype(font_path, font_size)

    # Draw the text on the image
    y_text = 0
    for line in wrapped_text:
        draw.text((50, y_text), line, font=font, fill="black")
        y_text += line_height

    
    # Save the image
    img.save("static/attachments/"+image_file)

def word_to_img(wfile,fid):
    # Create a Document object
    document = Document()
    # Load a Word DOCX file
    document.LoadFromFile("static/attachments/"+wfile)
    # Or load a Word DOC file
    #document.LoadFromFile("Sample.doc")

    # Convert the document to a list of image streams
    image_streams = document.SaveImageToStreams(ImageType.Bitmap)

    # Incremental counter
    i = 1

    # Save each image stream to a PNG file
    for image in image_streams:
        image_name = "m"+str(fid)+"_"+str(i) + ".png"
        with open("static/attachments/"+image_name,'wb') as image_file:
            image_file.write(image.ToArray())
        i += 1

    # Close the document
    document.Close()
    return i



def emailsink(usermail,pwd,uname):

    mycursor = mydb.cursor()
    
    import email
    import imaplib
    mail = imaplib.IMAP4_SSL('imap.gmail.com')
    (retcode, capabilities) = mail.login(usermail,pwd)
    mail.list()
    mail.select('inbox')

    subj1="Spam-Spoiler"
    n=0
    j=0
    ff=""
    fu=""
    fname=""
    imgname=""
    mess=""
    sender=""
    subj=""
    
    (retcode, messages) = mail.search(None, '(UnSeen)')
    if retcode == 'OK':
       
       for num in messages[0].split() :
          print ('Processing ')
          n=n+1

          mycursor.execute("SELECT max(id)+1 FROM read_data")
          maxid = mycursor.fetchone()[0]
          if maxid is None:
              maxid=1
                
          if n<4:
              typ, data = mail.fetch(num,'(RFC822)')
              
              for response_part in data:
                 if isinstance(response_part, tuple):
                     original = email.message_from_bytes(response_part[1])
                     
                    # print (original['From'])
                    # print (original['Subject'])
                     raw_email = data[0][1]
                     raw_email_string = raw_email.decode('utf-8')
                     email_message = email.message_from_string(raw_email_string)
                     
                     for part in email_message.walk():
                                ##
                                
                                if part.get_content_maintype() == 'multipart':
                                    #print(part.as_string())
                                    print("a")
                                    fu="1"

                                    continue
                                   
                                if part.get('Content-Disposition') is None:
                                    #print(part.as_string())
                                    print("b")
                                    continue
                               
                                if fu=="1":
                                    fileName = part.get_filename()
                                    print('file names processed ...')
                                    print(fileName)
                                    j+=1
                                    f1=fileName.split(".")
                                    
                                    if f1[1]=="txt" or f1[1]=="docx" or f1[1]=="pdf":
                                        
                                        fname="f"+str(maxid)+"_"+fileName
                                        
                                        ff+=fname+"|"
                                        
                                        #if bool(fileName):
                                        #filePath = os.path.join("static/attachments", fname)
                                        #if not os.path.isfile(filePath):
                                        print(fname)
                                            
                                        fp = open("static/attachments/"+fname, 'wb')
                                        fp.write(part.get_payload(decode=True))
                                        fp.close()
                                        #print('fp closed ...')

                                        if (part.get_content_type() == "text/plain"): # ignore attachments/html
                                            body = part.get_payload(decode=True)
                                            '''save_string = str(r"data.txt" )
                                            myfile = open(save_string, 'a')
                                            myfile.write(original['From']+'\n')
                                            myfile.write(original['Subject']+'\n')            
                                            myfile.write(body.decode('utf-8'))'''
                                            subj=original['Subject']
                                            sender=original['From']
                                            mess=body.decode('utf-8')
                                            

                                        if f1[1]=="txt":
                                            fp1=open("static/testdata.txt","r")
                                            val=fp1.read()
                                            fp1.close()

                                            fp12=open("static/attachments/"+fname,"r")
                                            txt=fp12.read()
                                            fp12.close()
                                            sval=val.split("|")
                                            x=0
                                            for sval2 in sval:
                                                if sval2 in txt:
                                                    x+=1
                                            if x>0:
                                                text_to_image(fname,maxid)
                                                n1=1
                                                #Delete mail
                                                mail.store(num,'+FLAGS',r'(\Deleted)')

                                                mess1="Malicious mail has deleted ***"
                                                sendmail(usermail,mess1,maxid,n1)
                                                sender=original['From']
                                                subj=original['Subject']
                                                sql = "INSERT INTO read_data(id,subject,sender,uname,message,spam_st,filename,img_count) VALUES (%s, %s, %s, %s, %s, %s,%s,%s)"
                                                val = (maxid,subj,sender,uname,'','0',fname,n1)
                                                
                                                mycursor.execute(sql, val)
                                                mydb.commit()
                                                print("mail sent")
                                                
                                        elif f1[1]=="docx":
                                            doc = docx.Document("static/attachments/"+fname)
                                            all_paras = doc.paragraphs
                                           
                                            fp1=open("static/testdata.txt","r")
                                            val=fp1.read()
                                            fp1.close()
                                            sval=val.split("|")
                                            x=0
                                            for sval2 in sval:
                                                for para in all_paras:
                                                    #print(para.text)
                                                    txt=para.text
                                                    if sval2 in txt:
                                                        x+=1
                                            if x>0:
                                                nn=word_to_img(fname,maxid)
                                                n1=nn-1
                                                #Delete mail
                                                mail.store(num,'+FLAGS',r'(\Deleted)')

                                                mess1="Malicious mail has deleted ***"
                                                sendmail(usermail,mess1,maxid,n1)
                                                sender=original['From']
                                                subj=original['Subject']
                                                sql = "INSERT INTO read_data(id,subject,sender,uname,message,spam_st,filename,img_count) VALUES (%s, %s, %s, %s, %s, %s,%s,%s)"
                                                val = (maxid,subj,sender,uname,'','0',fname,n1)                                                
                                                mycursor.execute(sql, val)
                                                mydb.commit()
                                                print("mail sent")
                                                
                                            #print(document.paragraphs)
                                        elif f1[1]=="pdf":
                                            p1=fname.split(".")
                                            fname2=p1[0]+".docx"

                                            convert_pdf2docx("static/attachments/"+fname,"static/attachments/"+fname2)
                                            
                                            doc = docx.Document("static/attachments/"+fname2)
                                            all_paras = doc.paragraphs
                                         
                                            fp1=open("static/testdata.txt","r")
                                            val=fp1.read()
                                            fp1.close()
                                            sval=val.split("|")
                                            x=0
                                            for sval2 in sval:
                                                for para in all_paras:
                                                    #print(para.text)
                                                    txt=para.text
                                                    if sval2 in txt:
                                                        x+=1
                                            if x>0:
                                                nn=word_to_img(fname2,maxid)
                                                n1=nn-1
                                                #Delete mail
                                                mail.store(num,'+FLAGS',r'(\Deleted)')

                                                mess1="Malicious mail has deleted ***"
                                                sendmail(usermail,mess1,maxid,n1)
                                                
                                                sender=original['From']
                                                subj=original['Subject']
                                                sql = "INSERT INTO read_data(id,subject,sender,uname,message,spam_st,filename,img_count) VALUES (%s, %s, %s, %s, %s, %s,%s,%s)"
                                                val = (maxid,subj,sender,uname,'','0',fname,n1)
                                                mycursor.execute(sql, val)
                                                mydb.commit()
                                                print("mail sent")
                                                
                                ##
                                '''print("#####*******#####")
                                ff2=ff.split("|")
                                ff3=len(ff2)-1
                                ff4=ff3-1
                                f=0
                                while f<ff4:
                                    ff5=ff2[f]
                                    print(ff5)
                                    f+=1
                                ##'''        
                                '''mess=""
                                sender=""
                                subj=""
                                if (part.get_content_type() == "text/plain"): # ignore attachments/html
                                      body = part.get_payload(decode=True)
                                      save_string = str(r"data.txt" )
                                      myfile = open(save_string, 'a')
                                      myfile.write(original['From']+'\n')
                                      myfile.write(original['Subject']+'\n')            
                                      myfile.write(body.decode('utf-8'))
                                      subj=original['Subject']
                                      sender=original['From']
                                      mess=body.decode('utf-8')
                                      

                                      
                                      myfile.write('**********\n')
                                      myfile.close()'''

                              
                                
                                '''if subj1==subj:
                                          print("subject")

                                      else:

                                          ###
                                          x=0
                                          y=0
                                          spam_st=""
                                          f1=open("spammail.txt","r")
                                          dat=f1.read()
                                          f1.close()

                                          
                                          #tt1=t1.rstrip()
                                          #stm=t1.split(" ")
                                          #for sm in stm:
                                          #    if sm in dat:
                                          #        y+=1
                                          #        print("yes1")
                                          #        break
                                          #    else:
                                          #        print("no1")
                                                      
                                          dat1=dat.split("|")
                                          for rd in dat1:
                                              rd1=rd.split('##')
                                              spam_st=rd1[1]
                                              t1=mess
                                              t2=rd1[0] #rd.strip()

                                              tt1=t1.rstrip()
                                              stm=tt1.split(" ")
                                              print(stm)
                                              for sm in stm:
                                                  if sm in t2:
                                                      y+=1
                                                      print("yes1")
                                                      break
                                                  else:
                                                      print("no1")
                                              ####     
                                              if t2 in t1:
                                                  x+=1
                                                  print("yes")
                                                  break
                                              else:
                                                  print("no")

                                              ###
                                              
                                              
                                          mail_det=""

                                          if x>0 or y>0:
                                              
                                              print(spam_st)
                                              if spam_st=="1":
                                                  mail_det="Fraudulent"
                                              elif spam_st=="2":
                                                  mail_det="Harrasment"
                                              elif spam_st=="3":
                                                  mail_det="Suspicious"
                                                  
                                              mycursor = mydb.cursor()
                                              mycursor.execute("SELECT max(id)+1 FROM read_data")
                                              maxid = mycursor.fetchone()[0]
                                              if maxid is None:
                                                  maxid=1
                                                    
                                              sql = "INSERT INTO read_data(id,subject,sender,uname,message,spam_st) VALUES (%s, %s, %s, %s, %s, %s)"
                                              val = (maxid,subj,sender,uname,mess,mail_det)
                                              mycursor.execute(sql, val)
                                              mydb.commit()

                                              ##
                                              #Reply mail
                                              
                                              mess1=mail_det+" mail has deleted *** "+mess+" *** "
                                              sendmail(usermail,mess1)

                                              #Delete mail
                                              mail.store(num,'+FLAGS',r'(\Deleted)')
                                              
                                          ###
                                else:
                                      continue'''

                 typ, data = mail.store(num,'+FLAGS','\\Seen')

    #print (n)
    return n
                 
    
@app.route('/userhome', methods=['GET', 'POST'])
def userhome():
   
    msg=""
    uname=""
    
    
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt","r")
    uname=ff.read()
    ff.close()

    mycursor = mydb.cursor()
    #uname="rajan"
    mycursor.execute("SELECT * FROM register where uname=%s",(uname,))
    det = mycursor.fetchone()
    em=det[3]
    email=det[3]
    pwd=det[6]

        
    return render_template('web/userhome.html',msg=msg,det=det)             

@app.route('/setting', methods=['GET', 'POST'])
def setting():
   
    msg=""
    uname=""
    act=request.args.get("act")
    
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt","r")
    uname=ff.read()
    ff.close()

    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM register where uname=%s",(uname,))
    det = mycursor.fetchone()
    em=det[3]
    pw=det[6]
        
    if request.method=='POST':
        
        email=request.form['email']
        password=request.form['pass']

        mycursor.execute("update register set email=%s,password=%s where uname=%s",(email,password,uname))
        mydb.commit()
       
        return redirect(url_for('setting',act='1'))
        
    return render_template('web/setting.html',msg=msg,em=em,pw=pw,act=act,det=det)



@app.route('/spam_detect', methods=['GET', 'POST'])
def spam_detect():
   
    msg=""
    uname=""
    
    
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt","r")
    uname=ff.read()
    ff.close()

    mycursor = mydb.cursor()

    #uname="rajan"
    mycursor.execute("SELECT * FROM register where uname=%s",(uname,))
    det = mycursor.fetchone()

    now = datetime.now()
    rdate=now.strftime("%d-%m-%Y")
    rtime=now.strftime("%H-M")
    dtt=rdate+" "+rtime
    #########

        
    return render_template('web/spam_detect.html',msg=msg,det=det)

@app.route('/page', methods=['GET', 'POST'])
def page():
   
    msg=""
    uname=""
    
    
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt","r")
    uname=ff.read()
    ff.close()

    mycursor = mydb.cursor()

    #uname="rajan"
    mycursor.execute("SELECT * FROM register where uname=%s",(uname,))
    det = mycursor.fetchone()
    em=det[3]
    email=det[3]
    pwd=det[6]
    print(email)
    print(pwd)

    now = datetime.now()
    rdate=now.strftime("%d-%m-%Y")
    rtime=now.strftime("%H-M")
    dtt=rdate+" "+rtime
    #########

    res=emailsink(email,pwd,uname)
    unread=res
    
    ########
    mycursor.execute("SELECT * FROM read_data where uname=%s order by id desc",(uname,))
    data = mycursor.fetchall()
    
        
    return render_template('web/page.html',msg=msg,em=em,data=data,det=det,unread=unread)


@app.route('/admin', methods=['GET', 'POST'])
def admin():   
    msg=""
    uname=""    
    
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt","r")
    uname=ff.read()
    ff.close()


    return render_template('web/admin.html',msg=msg)

@app.route('/load_data', methods=['GET', 'POST'])
def load_data():   
    msg=""
    uname=""    
    
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt","r")
    uname=ff.read()
    ff.close()

    act = request.args.get('act')

    pd.set_option("display.max_colwidth", 200) 
    warnings.filterwarnings("ignore") #ignore warnings

    #dataset/SEFACED_Email_Forensic_Dataset1.csv
    data = pd.read_csv(
        "static/dataset/spam.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    
    #dat1 = pd.read_csv("static/dataset/SEFACED_Email_Forensic_Dataset1.csv", header=0)
    #dat=dat1.head()
    data1=[]
    i=0
    for ds in data.values:
        if i<=200:
            data1.append(ds)
        i+=1

    return render_template('web/load_data.html',msg=msg,data1=data1)


#NLP
def lower(text):
    return text.lower()

def remove_specChar(text):
    return re.sub("#[A-Za-z0-9_]+", ' ', text)

def remove_link(text):
    return re.sub('@\S+|https?:\S+|http?:\S|[^A-Za-z0-9]+', ' ', text)

def remove_stopwords(text):
    return " ".join([word for word in 
                     str(text).split() if word not in STOPWORDS])

def stemming(text):
    return " ".join([stemmer.stem(word) for word in text.split()])

def stem_s(word):
    ww=word.split(" ")
    wd=[]
    
    for wr in ww:
        w1=len(wr)-1
        w2=len(wr)
        wrr=wr[w1:w2]
        if wrr == 's':
            wd.append(wr[:-1])
        else:
            wd.append(wr)
    res=" ".join(wd)
    return res
#def lemmatizer_words(text):
#    return " ".join([lematizer.lemmatize(word) for word in text.split()])

def cleanTxt(text):
    text = lower(text)
    text = remove_specChar(text)
    text = remove_link(text)
    text = remove_stopwords(text)
    #text = stemming(text)
    text = stem_s(text)
    
    return text


def generate_token(query):
    q4=[]
    query1=""
    q1=query.split(" ")
    for q11 in q1:
        q2=q11.split(".")
        q3="".join(q2)
        q4.append(q3)
    query1=" ".join(q4)
        
    text=cleanTxt(query1)
    stop_words = ['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves','The', 'you', "you're", "you've", "you'll", "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', "wouldn't",',','.','I','\'','-','/']
    ##Tokenize

    
    #doc = nlp(text)
    doc=text.split(" ")
    text_tokens=text.split(" ")
    
    #text_tokens = [token.text for token in text]
    ##text_tokens=tokenize_by_word(text)
    ##text_tokens =word_tokenize(text)
    tokens_without_sw = [word for word in text_tokens if not word in stop_words]
    return tokens_without_sw

@app.route('/preprocess', methods=['GET', 'POST'])
def preprocess():
    data1=[]
    data2=[]
    data3=[]

    df = pd.read_csv(
        "static/dataset/spam.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )

    i=0
    for ds in df.values:
        dt1=[]
        
        if i<200:
            if ds[1]=="":
                s=1
            else:
                dt1.append(ds[1])
                dd=generate_token(ds[1])
                dt1.append(dd)
                data1.append(dt1)        
        i+=1

    i=0
    for ds2 in df.values:
        dt2=[]
        
        if i<200:
            if ds2[1]=="":
                s=1
            else:
                dt2.append(ds2[1])
                dd2=cleanTxt(ds2[1])
                dt2.append(dd2)
                data2.append(dt2)        
        i+=1
    
    return render_template('web/preprocess.html',data1=data1,data2=data2,data3=data3)

def plot_ngrams(ax, texts, ngram_range=(2, 2), num_top_ngrams=25, title=''):
    # Initialize count vectorizer
    vectorizer = CountVectorizer(ngram_range=ngram_range)
    
    # Fit and transform the texts
    X = vectorizer.fit_transform(texts)
    
    # Get feature names
    feature_names = vectorizer.get_feature_names_out()
    
    # Sum the occurrences of each n-gram
    ngram_counts = X.sum(axis=0).A1
    
    # Create a dictionary of n-grams and their counts
    ngram_dict = dict(zip(feature_names, ngram_counts))
    
    # Sort the dictionary by counts in descending order
    sorted_ngrams = sorted(ngram_dict.items(), key=lambda x: x[1], reverse=True)
    
    # Select top N n-grams
    top_ngrams = sorted_ngrams[:num_top_ngrams]
    
    # Plot the top N n-grams
    sns.barplot(ax=ax, x=[ngram[1] for ngram in top_ngrams],
                y=[ngram[0] for ngram in top_ngrams],
                orient="h",
                width=0.5,
                palette='Spectral')
    ax.set_xlabel('Frequency')
    ax.set_ylabel('N-gram')
    ax.set_title(title)
    
@app.route('/feature', methods=['GET', 'POST'])
def feature():
    data1=[]
  
    '''emails = pd.read_csv(
        "static/dataset/spam.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )'''

    emails= pd.read_csv("static/dataset/spam.csv", encoding="ISO-8859-1")

    emails.shape
    emails.head()
    emails = emails.rename(columns={"v1": "Classify", "v2": "Email"})
    print(emails.isnull().sum())
    print("Total Duplicates:", emails.duplicated().sum())
    emails = emails.drop_duplicates()
    print("Total Duplicates:", emails.duplicated().sum())

    total_messages = len(emails)
    print("Total number of messages:", total_messages)

    '''spam_count = len(emails[emails['Classify'] == 'spam'])
    ham_count = len(emails[emails['Classify'] == 'ham'])

    total_messages = len(emails)
    spam_percentage = (spam_count / total_messages) * 100
    ham_percentage = (ham_count / total_messages) * 100

    # Create a pie chart
    labels = ['Spam', 'Regular (ham)']
    sizes = [spam_percentage, ham_percentage]
    colors = ['#ff9999', '#66b3ff']
    explode = (0.1, 0)  # Explode the first slice (spam)'''
    stop_words = ['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves','The', 'you', "you're", "you've", "you'll", "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', "wouldn't",',','.','I','\'','-','/']
    
    # Remove punctuation marks
    emails['Email'] = emails['Email'].apply(lambda x: x.translate(str.maketrans('', '', string.punctuation)))

    # Convert email text to lowercase
    emails['Email'] = emails['Email'].str.lower()

    # Remove stopwords
    #stop_words = set(stopwords.words('english'))
    emails['Email'] = emails['Email'].apply(lambda x: ' '.join([word for word in x.split() if word not in stop_words]))

    # Split each email into individual words
    all_words = ' '.join(emails['Email']).split()

    # Count the frequency of each word
    word_counts = Counter(all_words)

    # Get the most common words
    most_common_words = word_counts.most_common(100)

    print (most_common_words[:10])

    # Create a word cloud object
    wordcloud = WordCloud(width=800, height=400, background_color='white')

    # Generate the word cloud from the most_common_words list
    wordcloud.generate_from_frequencies(dict(most_common_words))

    # Plot the word cloud
    plt.figure(figsize=(10, 6))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    #plt.savefig('static/graph/ff.png')
    plt.close()
    #plt.show()


    # Filter regular messages (ham)
    regular_messages = emails[emails['Classify'] == 'ham']

    # Count the number of regular messages containing the word "hello"
    count_hello = regular_messages['Email'].str.contains('hello', case=False).sum()

    # Print the result
    print("Number of regular messages containing the word 'hello':", count_hello)


    import re

    phone_number_count = 0

    for message in emails.loc[emails['Classify'] == 'spam', 'Email']:
        if re.search(r'\b\d{10}\b', message):  # assuming phone numbers are 10 digits long
            phone_number_count += 1

    print("Number of spam messages containing a phone number:", phone_number_count)
    word_counts = emails['Email'].apply(lambda x: len(str(x).split()))
    average_word_count = word_counts.mean()

    emails['number_of_characters_in_the_message'] = emails['Email'].apply(len)

    emails['number_of_characters_in_the_message'].hist(bins=50)
    plt.xlabel('Number of characters in the message')
    plt.ylabel('Frequency')
    plt.title('Distribution of number of characters in messages')
    #plt.savefig('static/graph/ff2.png')
    #plt.show()

    ##
    # Filter spam and non-spam messages
    '''spam_texts = emails[emails['Classify'] == 'spam']['Email']
    non_spam_texts = emails[emails['Classify'] == 'ham']['Email']

    #Visualization
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    plot_ngrams(axes[0], spam_texts, title='Top Bigrams in Spam Messages')
    plot_ngrams(axes[1], non_spam_texts, title='Top Bigrams in Non-Spam Messages')
    axes[0].grid(axis='x')
    axes[1].grid(axis='x')
    plt.tight_layout()
    plt.show()
    ###
    vectorizer = CountVectorizer()

    # Bag of words
    bow_text = vectorizer.fit_transform(emails["Email"])

    # Fetch the vocabulary set
    #print(f"10 Bag Of Words Features: {vectorizer.get_feature_names_out()[100:110]}")
    #print(f"Total number of vocab words: {len(vectorizer.vocabulary_)}")
    transformed_bow = vectorizer.transform(emails["Email"])

    # TF-IDF
    tfidf_transformer = TfidfTransformer().fit(transformed_bow)
    text_tfidf = tfidf_transformer.transform(transformed_bow)'''




    return render_template('web/feature.html',data1=data1)

@app.route('/classify', methods=['GET', 'POST'])
def classify():
    data1=[]
  
    df = pd.read_csv(
        "static/dataset/spam.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    data1=[]
    i=0
    for ds in df.values:
        if i<=200:
            data1.append(ds)
        i+=1

    x=0
    y=0
    for ds1 in df.values:
        if ds1[0]=="ham":
            x+=1
        if ds1[0]=="spam":
            y+=1
        
    cname=['Not Spam','Spam']
    dd2=[x,y]
    gt=x+10
    ##
    doc = cname #list(data.keys())
    values = dd2 #list(data.values())
    
    print(doc)
    print(values)
    fig = plt.figure(figsize = (10, 8))
     
    # creating the bar plot
    cc=['blue','brown']
    plt.bar(doc, values, color =cc,
            width = 0.4)
 

    plt.ylim((1,gt))
    plt.xlabel("Class")
    plt.ylabel("Count")
    plt.title("")

    rr=randint(100,999)
    fn="tclass.png"
    #plt.xticks(rotation=20)
    plt.savefig('static/graph/'+fn)
    
    plt.close()
    ##

    return render_template('web/classify.html',data1=data1)
    

@app.route('/train_data', methods=['GET', 'POST'])
def train_data():
    msg=""
    act = request.args.get('act')

    pd.set_option("display.max_colwidth", 200) 
    warnings.filterwarnings("ignore") #ignore warnings

    #dataset/SEFACED_Email_Forensic_Dataset1.csv
    data = pd.read_csv(
        "static/dataset/train.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    #dat1 = pd.read_csv("static/dataset/SEFACED_Email_Forensic_Dataset1.csv", header=0)
    #dat=dat1.head()
    data1=[]
    i=0
    for ds in data.values:
        #if i<=200:
        data1.append(ds)
        #i+=1
    '''plt.rc("axes.spines", right=False, top=False)
    plt.rc("font", family="serif")
            
    data = pd.read_csv(
        "static/dataset/data1.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    dat=data.head()
    data1=[]
    for ds in dat.values:
        data1.append(ds)'''

    '''for label, cmap in zip(["ham", "spam"], ["winter", "autumn"]):
        text = data.query("label == @label")["text"].str.cat(sep=" ")
        plt.figure(figsize=(10, 6))
        #wc = WordCloud(width=1000, height=600, background_color="#f8f8f8", colormap=cmap)
        #wc.generate_from_text(text)
        #plt.imshow(wc)
        #plt.axis("off")
        #plt.title(f"Words Commonly Used in ${label}$ Messages", size=20)
        #plt.show()'''

    '''data["length (words)"] = data["text"].str.split().apply(len)
    print(data.groupby("label").agg([min, max, "mean"]))
    ax = data.boxplot(by="label", figsize=(6, 4.5))
    _ = ax.set_title("")
    #plt.show()
    #plt.close()
    ######

    X_train, X_test, y_train, y_test = train_test_split(
        data["text"], data["label"], random_state=8, stratify=data["label"]
    )

    _ = y_train.value_counts().plot.bar(
        color=["aqua", "orangered"], edgecolor="#555", alpha=0.5
    )
    #plt.show()
    #plt.close()'''
    #############
    

    return render_template('train_data.html',msg=msg,data1=data1)



@app.route('/process2', methods=['GET', 'POST'])
def process2():

    ##########
    plt.rc("axes.spines", right=False, top=False)
    plt.rc("font", family="serif")
    data = pd.read_csv(
        "static/dataset/data1.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    dat=data.head()

    

    data["length (words)"] = data["text"].str.split().apply(len)
    dataval=data.groupby("label").agg([min, max, "mean"])
    print(data.groupby("label").agg([min, max, "mean"]))
    ax = data.boxplot(by="label", figsize=(6, 4.5))
    _ = ax.set_title("")
    #plt.show()
    plt.savefig("static/dataset/graph2.png")
    plt.close()
    ######

    X_train, X_test, y_train, y_test = train_test_split(
        data["text"], data["label"], random_state=8, stratify=data["label"]
    )

    _ = y_train.value_counts().plot.bar(
        color=["aqua", "orangered"], edgecolor="#555", alpha=0.5
    )
    #plt.show()
    plt.savefig("static/dataset/graph3.png")
    plt.close()
    ############################################################################

    data = pd.read_csv(
        "static/dataset/data2.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    dat=data.head()

    

    data["length (words)"] = data["text"].str.split().apply(len)
    dataval2=data.groupby("label").agg([min, max, "mean"])
    print(data.groupby("label").agg([min, max, "mean"]))
    ax = data.boxplot(by="label", figsize=(6, 4.5))
    _ = ax.set_title("")
    #plt.show()
    plt.savefig("static/dataset/graph4.png")
    plt.close()
    ######

    X_train, X_test, y_train, y_test = train_test_split(
        data["text"], data["label"], random_state=8, stratify=data["label"]
    )

    _ = y_train.value_counts().plot.bar(
        color=["aqua", "orangered"], edgecolor="#555", alpha=0.5
    )
    #plt.show()
    plt.savefig("static/dataset/graph5.png")
    plt.close()
    #############################################################################
    data = pd.read_csv(
        "static/dataset/data3.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    dat=data.head()

    

    data["length (words)"] = data["text"].str.split().apply(len)
    dataval3=data.groupby("label").agg([min, max, "mean"])
    print(data.groupby("label").agg([min, max, "mean"]))
    ax = data.boxplot(by="label", figsize=(6, 4.5))
    _ = ax.set_title("")
    #plt.show()
    plt.savefig("static/dataset/graph6.png")
    plt.close()
    ######

    X_train, X_test, y_train, y_test = train_test_split(
        data["text"], data["label"], random_state=8, stratify=data["label"]
    )

    _ = y_train.value_counts().plot.bar(
        color=["aqua", "orangered"], edgecolor="#555", alpha=0.5
    )
    #plt.show()
    plt.savefig("static/dataset/graph7.png")
    plt.close()
    #################################################################
    
    return render_template('process2.html',dataval=dataval,dataval2=dataval2,dataval3=dataval3)

@app.route('/process3', methods=['GET', 'POST'])
def process3():

    df = pd.read_csv('static/dataset/spam11.csv')

    print(df.shape)
    dat=df.head()
    data1=[]
    for ds1 in dat.values:
        data1.append(ds1)
    
    dat2=df.describe()
    data2=[]
    drr=['count','mean','std','min','25%','50%','75%','max']
    i=0
    for ds2 in dat2.values:
        dt=[]
        dt.append(drr[i])
        dt.append(ds2)
        i+=1
        data2.append(dt)
    
    #df.info()
    dat3=df.corr()
    data3=[]
    for ds3 in dat3.values:
        data3.append(ds3)


    #visualize correlation of variable using pearson correlation
    plt.figure(figsize = (8,6))
    sns.heatmap(df.corr(), vmax = 0.9, cmap = 'YlGnBu')
    plt.title('Pearson Correlation', fontsize = 15, pad = 12, color = 'r')
    plt.savefig("static/dataset/ff_g1.png")
    #plt.show()

    #transform spam column to categorical data
    df.spam[df['spam'] == 0] = 'ham'
    df.spam[df['spam'] == 1] = 'spam'
    dat4=df.head()
    data4=[]
    for ds4 in dat4.values:
        data4.append(ds4)
    
    #analyze of spam status based on capital run length average
    dat5=pd.pivot_table(df, index = 'spam', values = 'capital_run_length_average', 
                   aggfunc = {'capital_run_length_average' : np.mean}).sort_values('capital_run_length_average', ascending = False)

    print(dat5)

    #analyze of spam status based on count of capital run length longest
    pd.pivot_table(df, index = 'spam', values = 'capital_run_length_longest',
                  aggfunc = {'capital_run_length_longest' : np.sum}).sort_values('capital_run_length_longest', ascending = False)


    #anayze of spam status based on count of capital run length total
    pd.pivot_table(df, index = 'spam', values = 'capital_run_length_total',
                  aggfunc = {'capital_run_length_total' : np.sum}).sort_values('capital_run_length_total', ascending = False)


    #anayze of spam status based on capital run length average, capital run length longest and capital run length total
    pd.pivot_table(df, index = 'spam', values = ['capital_run_length_average', 'capital_run_length_longest', 
                                                 'capital_run_length_total'], 
                   aggfunc = {'capital_run_length_average' : np.mean, 'capital_run_length_longest' : np.sum, 
                              'capital_run_length_total' : np.sum}).sort_values(['capital_run_length_average', 
                                                                                 'capital_run_length_longest', 
                                                                                 'capital_run_length_total'], ascending = False)


    #visualize the factor of spam message based on capital run length average, capital run length longest and capital run length total
    plt.figure(figsize = (14,6))
    chart = df.boxplot()
    chart.set_xticklabels(chart.get_xticklabels(), rotation = 90)
    plt.title('The Factor of Spam Message', fontsize = 15, pad = 12, color = 'b')
    plt.xlabel('Factor')
    plt.ylabel('Count')
    plt.savefig("static/dataset/ff_g2.png")
    #plt.show()


    return render_template('process3.html',data1=data1,data2=data2,data3=data3,data4=data4,dat5=dat5)

#RoBERTa--Malicious Content Detection
def RoBERTa():
    batch = collate_tokens(
        [roberta.encode(pair[0], pair[1]) for pair in batch_of_pairs], pad_idx=1
    )
    logprobs = roberta.predict('mnli', batch)
    print(logprobs.argmax(dim=1))
    label_map = {0: 'contradiction', 1: 'neutral', 2: 'entailment'}
    ncorrect, nsamples = 0, 0
    roberta.cuda()
    roberta.eval()
    with open('static/dataset/spam11.csv') as fin:
        fin.readline()
        for index, line in enumerate(fin):
            tokens = line.strip().split('\t')
            sent1, sent2, target = tokens[8], tokens[9], tokens[-1]
            tokens = roberta.encode(sent1, sent2)
            prediction = roberta.predict('mnli', tokens).argmax().item()
            prediction_label = label_map[prediction]
            ncorrect += int(prediction_label == target)
            nsamples += 1
    print('| Accuracy: ', float(ncorrect)/float(nsamples))

    def get_data():
        text = str(self.text[index])
        text = " ".join(text.split())

        inputs = self.tokenizer.encode_plus(
            text,
            None,
            add_special_tokens=True,
            max_length=self.max_len,
            pad_to_max_length=True,
            return_token_type_ids=True
        )
        ids = inputs['input_ids']
        token_type_ids = inputs["token_type_ids"]
        return {
            'ids': torch.tensor(ids, dtype=torch.long),
            'token_type_ids': torch.tensor(token_type_ids, dtype=torch.long),
            'targets': torch.tensor(self.targets[index], dtype=torch.float)
        }

    train_size = 0.8
    train_data=new_df.sample(frac=train_size,random_state=200)
    test_data=new_df.drop(train_data.index).reset_index(drop=True)
    train_data = train_data.reset_index(drop=True)


    print("FULL Dataset: {}".format(new_df.shape))
    print("TRAIN Dataset: {}".format(train_data.shape))
    print("TEST Dataset: {}".format(test_data.shape))

    training_set = SentimentData(train_data, tokenizer, MAX_LEN)
    testing_set = SentimentData(test_data, tokenizer, MAX_LEN)

    train_params = {'batch_size': TRAIN_BATCH_SIZE,
                'shuffle': True,
                'num_message': 0
                }

    test_params = {'batch_size': VALID_BATCH_SIZE,
                    'shuffle': True,
                    'num_message': 0
                    }

    training_loader = DataLoader(training_set, **train_params)
    testing_loader = DataLoader(testing_set, **test_params)

def train(epoch):
    tr_loss = 0
    n_correct = 0
    nb_tr_steps = 0
    nb_tr_examples = 0
    model.train()
    for _,data in tqdm(enumerate(training_loader, 0)):
        ids = data['ids'].to(device, dtype = torch.long)
        mask = data['mask'].to(device, dtype = torch.long)
        token_type_ids = data['token_type_ids'].to(device, dtype = torch.long)
        targets = data['targets'].to(device, dtype = torch.long)

        outputs = model(ids, mask, token_type_ids)
        loss = loss_function(outputs, targets)
        tr_loss += loss.item()
        big_val, big_idx = torch.max(outputs.data, dim=1)
        n_correct += calcuate_accuracy(big_idx, targets)

        nb_tr_steps += 1
        nb_tr_examples+=targets.size(0)
        
        if _%5000==0:
            loss_step = tr_loss/nb_tr_steps
            accu_step = (n_correct*100)/nb_tr_examples 
            print(f"Training Loss per 5000 steps: {loss_step}")
            print(f"Training Accuracy per 5000 steps: {accu_step}")

        optimizer.zero_grad()
        loss.backward()
        # # When using GPU
        optimizer.step()

    print(f'The Total Accuracy for Epoch {epoch}: {(n_correct*100)/nb_tr_examples}')
    epoch_loss = tr_loss/nb_tr_steps
    epoch_accu = (n_correct*100)/nb_tr_examples
    print(f"Training Loss Epoch: {epoch_loss}")
    print(f"Training Accuracy Epoch: {epoch_accu}")

    return 
def valid(model, testing_loader):
    model.eval()
    n_correct = 0; n_wrong = 0; total = 0; tr_loss=0; nb_tr_steps=0; nb_tr_examples=0
    with torch.no_grad():
        for _, data in tqdm(enumerate(testing_loader, 0)):
            ids = data['ids'].to(device, dtype = torch.long)
            mask = data['mask'].to(device, dtype = torch.long)
            token_type_ids = data['token_type_ids'].to(device, dtype=torch.long)
            targets = data['targets'].to(device, dtype = torch.long)
            outputs = model(ids, mask, token_type_ids).squeeze()
            loss = loss_function(outputs, targets)
            tr_loss += loss.item()
            big_val, big_idx = torch.max(outputs.data, dim=1)
            n_correct += calcuate_accuracy(big_idx, targets)

            nb_tr_steps += 1
            nb_tr_examples+=targets.size(0)
            
            if _%5000==0:
                loss_step = tr_loss/nb_tr_steps
                accu_step = (n_correct*100)/nb_tr_examples
                print(f"Validation Loss per 100 steps: {loss_step}")
                print(f"Validation Accuracy per 100 steps: {accu_step}")
    epoch_loss = tr_loss/nb_tr_steps
    epoch_accu = (n_correct*100)/nb_tr_examples
    print(f"Validation Loss Epoch: {epoch_loss}")
    print(f"Validation Accuracy Epoch: {epoch_accu}")    
    return epoch_accu
#######################

    
@app.route('/process4', methods=['GET', 'POST'])
def process4():

    ##########
    plt.rc("axes.spines", right=False, top=False)
    plt.rc("font", family="serif")
    data = pd.read_csv(
        "static/dataset/data1.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    dat=data.head()

    

    data["length (words)"] = data["text"].str.split().apply(len)
    dataval=data.groupby("label").agg([min, max, "mean"])
    print(data.groupby("label").agg([min, max, "mean"]))
    ax = data.boxplot(by="label", figsize=(6, 4.5))
    _ = ax.set_title("")
    #plt.show()
    #plt.savefig("static/dataset/graph2.png")
    #plt.close()
    ######

    X_train, X_test, y_train, y_test = train_test_split(
        data["text"], data["label"], random_state=8, stratify=data["label"]
    )

    #_ = y_train.value_counts().plot.bar(
    #    color=["aqua", "orangered"], edgecolor="#555", alpha=0.5
    #)
    #plt.show()
    #plt.xticks(rotation=0)
    #plt.savefig("static/dataset/graph3.png")
    #plt.close()
    ############################################################################

    data = pd.read_csv(
        "static/dataset/data2.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    dat=data.head()

    

    data["length (words)"] = data["text"].str.split().apply(len)
    dataval2=data.groupby("label").agg([min, max, "mean"])
    print(data.groupby("label").agg([min, max, "mean"]))
    ax = data.boxplot(by="label", figsize=(6, 4.5))
    _ = ax.set_title("")
    #plt.show()
    #plt.savefig("static/dataset/graph4.png")
    #plt.close()
    ######

    X_train, X_test, y_train, y_test = train_test_split(
        data["text"], data["label"], random_state=8, stratify=data["label"]
    )

    #_ = y_train.value_counts().plot.bar(
    #    color=["aqua", "orangered"], edgecolor="#555", alpha=0.5
    #)
    #plt.show()
    #plt.xticks(rotation=0)
    #plt.savefig("static/dataset/graph5.png")
    #plt.close()
    #############################################################################
    data = pd.read_csv(
        "static/dataset/data3.csv",
        header=0,
        encoding="latin-1",
        usecols=[0, 1],
        names=["label", "text"],
    )
    dat=data.head()

    

    data["length (words)"] = data["text"].str.split().apply(len)
    dataval3=data.groupby("label").agg([min, max, "mean"])
    print(data.groupby("label").agg([min, max, "mean"]))
    ax = data.boxplot(by="label", figsize=(6, 4.5))
    _ = ax.set_title("")
    #plt.show()
    #plt.savefig("static/dataset/graph6.png")
    #plt.close()
    ######

    X_train, X_test, y_train, y_test = train_test_split(
        data["text"], data["label"], random_state=8, stratify=data["label"]
    )

    #_ = y_train.value_counts().plot.bar(
    #    color=["aqua", "orangered"], edgecolor="#555", alpha=0.5
    #)
    #plt.show()
    #plt.xticks(rotation=0)
    #plt.savefig("static/dataset/graph7.png")
    #plt.close()
    #################################################################

    
    df = pd.read_csv('static/dataset/SEFACED_Email_Forensic_Dataset1.csv',delimiter=',',encoding='latin-1')
    df = df[['Class_Label','Text']]
    df = df[pd.notnull(df['Text'])]
    df.rename(columns = {'Message':'Text'}, inplace = True)
    print(df.head())
    data1=[]
    ##for ds in data.values:
        
    ##    data1.append(ds)

        
    dsf=df.shape
    print(dsf)
    print(dsf[0])


    df.index = range(dsf[0])
    df['Text'].apply(lambda x: len(x.split(' '))).sum()

    #cnt_pro = df['Class_Label'].value_counts()
    #plt.figure(figsize=(12,4))
    #sns.barplot(cnt_pro.index, cnt_pro.values, alpha=0.8)
    #plt.ylabel('Number of Occurrences', fontsize=12)
    #plt.xlabel('Class_Label', fontsize=12)
    #plt.xticks(rotation=90)
    #plt.savefig("static/dataset/graph1.png")
    #plt.show();


    
    return render_template('process4.html',dataval=dataval,dataval2=dataval2,dataval3=dataval3)



##########################
@app.route('/logout')
def logout():
    # remove the username from the session if it is there
    session.pop('username', None)
    return redirect(url_for('index'))



if __name__ == '__main__':
    app.secret_key = os.urandom(12)
    app.run(debug=True,host='0.0.0.0', port=5000)



